"""Тесты рендеринга резюме в документы."""

from datetime import UTC, datetime
from inspect import signature
from io import BytesIO
from types import SimpleNamespace
from typing import Any
from uuid import uuid4

import pytest
from docx import Document

from app.core.exceptions import ServiceUnavailableException
from app.dto.resumes import ResumeDTO, ResumeSectionDTO, SectionType
from app.services.resume_document import ResumeDocumentRenderer

TIMESTAMP = datetime(2026, 1, 1, tzinfo=UTC)


def resume_section(
    position: int,
    section_type: SectionType,
    content: dict[str, Any],
) -> ResumeSectionDTO:
    """Создаёт DTO секции для рендеринга."""
    return ResumeSectionDTO(
        id=uuid4(),
        resume_id=uuid4(),
        section_type=section_type,
        position=position,
        content=content,
        created_at=TIMESTAMP,
        updated_at=TIMESTAMP,
    )


def resume() -> ResumeDTO:
    """Возвращает резюме для рендеринга."""
    return ResumeDTO(
        id=uuid4(),
        user_id=uuid4(),
        title="Python Developer",
        created_at=TIMESTAMP,
        updated_at=TIMESTAMP,
        sections=(
            resume_section(0, SectionType.SUMMARY, {"text": "Backend developer"}),
            resume_section(
                1,
                SectionType.SKILLS,
                {"skills": [{"name": "Python", "level": "Senior"}]},
            ),
        ),
    )


def full_resume() -> ResumeDTO:
    """Возвращает резюме со всеми секциями в пользовательском порядке."""
    return ResumeDTO(
        id=uuid4(),
        user_id=uuid4(),
        title="Python <Developer>",
        created_at=TIMESTAMP,
        updated_at=TIMESTAMP,
        sections=(
            resume_section(
                3,
                SectionType.EXPERIENCE,
                {
                    "experiences": [
                        {
                            "company": "Acme & Co",
                            "position": "Backend Developer",
                            "start_date": "2024-01-01",
                            "end_date": None,
                            "description": "Built APIs",
                        }
                    ]
                },
            ),
            resume_section(1, SectionType.SUMMARY, {"text": "ATS-friendly summary"}),
            resume_section(
                2,
                SectionType.EDUCATION,
                {
                    "education": [
                        {
                            "institution": "State University",
                            "degree": "Bachelor",
                            "field": "Computer Science",
                            "start_date": "2019-09-01",
                            "end_date": "2023-06-30",
                        }
                    ]
                },
            ),
            resume_section(
                4,
                SectionType.SKILLS,
                {"skills": [{"name": "Python", "level": "Senior"}]},
            ),
            resume_section(
                5,
                SectionType.PROJECTS,
                {
                    "projects": [
                        {
                            "name": "AI Resume",
                            "description": "Resume builder",
                            "url": "https://example.com/resume",
                        }
                    ]
                },
            ),
            resume_section(
                6,
                SectionType.LANGUAGES,
                {"languages": [{"name": "English", "level": "C1"}]},
            ),
        ),
    )


def test_renders_docx_with_resume_content() -> None:
    """Рендерит DOCX с заголовком и секциями."""
    content = ResumeDocumentRenderer().render_docx(resume())
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Python Developer" in text
    assert "Backend developer" in text
    assert "Python" in text


def test_renderer_public_methods_accept_resume_dto() -> None:
    """Принимает DTO резюме на публичной границе рендерера."""
    renderer = ResumeDocumentRenderer()

    assert signature(renderer.render_docx).parameters["resume"].annotation is ResumeDTO
    assert signature(renderer.render_pdf).parameters["resume"].annotation is ResumeDTO


def test_renders_all_sections_in_position_order_for_docx() -> None:
    """Рендерит все секции DOCX в порядке позиции и без пустых полей."""
    content = ResumeDocumentRenderer().render_docx(full_resume())
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    expected_parts = [
        "Python <Developer>",
        "Summary",
        "Education",
        "Experience",
        "Skills",
        "Projects",
        "Languages",
        "2019-09-01 — 2023-06-30",
        "2024-01-01 — Present",
        "https://example.com/resume",
    ]
    indices = [text.index(part) for part in expected_parts[:7]]

    assert all(part in text for part in expected_parts)
    assert indices == sorted(indices)
    assert "None" not in text


def test_renders_all_sections_as_escaped_html_in_position_order() -> None:
    """Строит экранированный HTML со всеми секциями в порядке позиции."""
    html = ResumeDocumentRenderer()._render_html(full_resume())

    expected_headings = [
        "Summary",
        "Education",
        "Experience",
        "Skills",
        "Projects",
        "Languages",
    ]
    indices = [html.index(heading) for heading in expected_headings]

    assert "Python &lt;Developer&gt;" in html
    assert "Acme &amp; Co" in html
    assert "2024-01-01 — Present" in html
    assert indices == sorted(indices)


def test_renders_all_sections_for_pdf(monkeypatch: pytest.MonkeyPatch) -> None:
    """Передаёт в PDF-движок HTML со всеми секциями резюме."""
    rendered_html: list[str] = []

    class Html:
        """Имитирует HTML-объект WeasyPrint."""

        def __init__(self, string: str) -> None:
            """Сохраняет HTML для проверки."""
            self.string = string
            rendered_html.append(string)

        def write_pdf(self) -> bytes:
            """Возвращает содержимое тестового PDF."""
            return b"pdf"

    original_import = __import__

    def import_weasyprint(
        name: str,
        *args: object,
        **kwargs: object,
    ) -> object:
        """Подменяет WeasyPrint тестовой реализацией."""
        if name == "weasyprint":
            return SimpleNamespace(HTML=Html)
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", import_weasyprint)

    assert ResumeDocumentRenderer().render_pdf(full_resume()) == b"pdf"
    html = rendered_html[0]
    assert all(
        heading in html
        for heading in (
            "Summary",
            "Education",
            "Experience",
            "Skills",
            "Projects",
            "Languages",
        )
    )


def test_pdf_render_translates_missing_system_library_to_service_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Не раскрывает ошибку загрузки системной библиотеки WeasyPrint."""
    original_import = __import__

    def raise_weasyprint_import_error(
        name: str,
        *args: object,
        **kwargs: object,
    ) -> object:
        """Имитирует недоступность нативной зависимости WeasyPrint."""
        if name == "weasyprint":
            raise OSError("cannot load library 'libgobject-2.0-0'")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", raise_weasyprint_import_error)

    with pytest.raises(
        ServiceUnavailableException,
        match="PDF export is temporarily unavailable",
    ):
        ResumeDocumentRenderer().render_pdf(resume())
