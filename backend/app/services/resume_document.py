"""Рендерит резюме в ATS-friendly PDF и DOCX."""

from dataclasses import dataclass
from html import escape
from io import BytesIO
from typing import Any, cast

from docx import Document
from docx.shared import Pt

from app.core.exceptions import ServiceUnavailableException
from app.dto.resumes import ResumeDTO

PDF_EXPORT_UNAVAILABLE_MESSAGE = "PDF export is temporarily unavailable"


@dataclass(frozen=True)
class RenderedItem:
    """Представляет элемент секции для вывода в документ."""

    title: str | None
    details: tuple[str, ...]


@dataclass(frozen=True)
class RenderedSection:
    """Представляет секцию для вывода в документ."""

    heading: str
    items: tuple[RenderedItem, ...]


class ResumeDocumentRenderer:
    """Рендерит единое представление резюме в поддерживаемые форматы."""

    def render_pdf(self, resume: ResumeDTO) -> bytes:
        """Возвращает PDF-представление резюме."""
        try:
            from weasyprint import HTML  # type: ignore[import-untyped]

            return cast(bytes, HTML(string=self._render_html(resume)).write_pdf())
        except OSError as exc:
            raise ServiceUnavailableException(PDF_EXPORT_UNAVAILABLE_MESSAGE) from exc

    def render_docx(self, resume: ResumeDTO) -> bytes:
        """Возвращает DOCX-представление резюме."""
        document = Document()
        document.styles["Normal"].font.name = "Arial"
        document.styles["Normal"].font.size = Pt(10)
        document.add_heading(str(resume.title), level=0)

        for section in self._build_sections(resume):
            document.add_heading(section.heading, level=1)
            for item in section.items:
                if item.title:
                    document.add_heading(item.title, level=2)
                for detail in item.details:
                    document.add_paragraph(detail)

        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    def _render_html(self, resume: ResumeDTO) -> str:
        """Строит экранированный HTML из единого представления резюме."""
        sections = "".join(
            self._render_html_section(section)
            for section in self._build_sections(resume)
        )
        return (
            "<html><head><meta charset='utf-8'><style>"
            "body{font-family:Arial,sans-serif;font-size:11pt;color:#111}"
            "h1{font-size:22pt}h2{font-size:15pt;margin-top:18pt}"
            "h3{font-size:12pt;margin-bottom:0}</style></head><body>"
            f"<h1>{escape(str(resume.title))}</h1>{sections}</body></html>"
        )

    def _build_sections(self, resume: ResumeDTO) -> tuple[RenderedSection, ...]:
        """Преобразует секции резюме в единое представление по их позиции."""
        sections = sorted(
            resume.sections,
            key=lambda section: section.position,
        )
        return tuple(
            rendered_section
            for section in sections
            if (
                rendered_section := self._build_section(
                    str(section.section_type), section.content
                )
            )
            is not None
        )

    def _build_section(
        self,
        section_type: str,
        content: dict[str, Any],
    ) -> RenderedSection | None:
        """Создаёт представление одной известной секции резюме."""
        if section_type == "summary":
            return RenderedSection(
                heading="Summary",
                items=(RenderedItem(None, (content["text"],)),),
            )
        if section_type == "experience":
            return RenderedSection(
                heading="Experience",
                items=tuple(
                    RenderedItem(
                        title=f"{item['position']} — {item['company']}",
                        details=self._non_empty(
                            self._period(item), item.get("description")
                        ),
                    )
                    for item in content["experiences"]
                ),
            )
        if section_type == "education":
            return RenderedSection(
                heading="Education",
                items=tuple(
                    RenderedItem(
                        title=(
                            f"{item['degree']} in {item['field']} — "
                            f"{item['institution']}"
                        ),
                        details=(self._period(item),),
                    )
                    for item in content["education"]
                ),
            )
        if section_type == "skills":
            return RenderedSection(
                heading="Skills",
                items=(
                    RenderedItem(
                        None,
                        tuple(self._skill_label(item) for item in content["skills"]),
                    ),
                ),
            )
        if section_type == "projects":
            return RenderedSection(
                heading="Projects",
                items=tuple(
                    RenderedItem(
                        title=item["name"],
                        details=self._non_empty(
                            item.get("description"), item.get("url")
                        ),
                    )
                    for item in content["projects"]
                ),
            )
        if section_type == "languages":
            return RenderedSection(
                heading="Languages",
                items=(
                    RenderedItem(
                        None,
                        tuple(
                            f"{item['name']} ({item['level']})"
                            for item in content["languages"]
                        ),
                    ),
                ),
            )
        return None

    @staticmethod
    def _render_html_section(section: RenderedSection) -> str:
        """Преобразует единое представление секции в HTML."""
        items = "".join(
            "<article>"
            f"{f'<h3>{escape(item.title)}</h3>' if item.title else ''}"
            f"{''.join(f'<p>{escape(detail)}</p>' for detail in item.details)}"
            "</article>"
            for item in section.items
        )
        return f"<section><h2>{escape(section.heading)}</h2>{items}</section>"

    @staticmethod
    def _period(item: dict[str, Any]) -> str:
        """Форматирует период с ISO-датами и значением для текущего периода."""
        end_date = item.get("end_date") or "Present"
        return f"{item['start_date']} — {end_date}"

    @staticmethod
    def _skill_label(item: dict[str, Any]) -> str:
        """Форматирует навык с необязательным уровнем."""
        level = item.get("level")
        return f"{item['name']} ({level})" if level else item["name"]

    @staticmethod
    def _non_empty(*values: str | None) -> tuple[str, ...]:
        """Возвращает только непустые значения для вывода."""
        return tuple(value for value in values if value)
