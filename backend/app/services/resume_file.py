"""Проверяет файлы резюме и извлекает из них текст."""

from io import BytesIO
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader

from app.core.exceptions import ValidationException

PDF_SIGNATURE = b"%PDF-"
ZIP_SIGNATURE = b"PK\x03\x04"


class ResumeFileExtractor:
    """Проверяет и читает текстовые PDF и DOCX."""

    def __init__(self, max_file_size: int) -> None:
        """Инициализирует экземпляр."""
        self.max_file_size = max_file_size

    def extract(self, filename: str, content: bytes) -> str:
        """Возвращает текст из поддерживаемого файла резюме."""
        self._validate_size(content)
        if content.startswith(PDF_SIGNATURE):
            return self._extract_pdf(content)
        if content.startswith(ZIP_SIGNATURE):
            return self._extract_docx(content)
        raise ValidationException("Unsupported file format")

    def _validate_size(self, content: bytes) -> None:
        """Проверяет размер файла."""
        if not content:
            raise ValidationException("File is empty")
        if len(content) > self.max_file_size:
            raise ValidationException("File is too large")

    def _extract_pdf(self, content: bytes) -> str:
        """Извлекает текст из PDF."""
        try:
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as exc:
            raise ValidationException("Invalid PDF file") from exc
        if not text:
            raise ValidationException("PDF does not contain text")
        return text

    def _extract_docx(self, content: bytes) -> str:
        """Извлекает текст из DOCX."""
        try:
            with ZipFile(BytesIO(content)) as archive:
                names = set(archive.namelist())
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValidationException("Unsupported file format")
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            paragraphs.extend(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
            text = "\n".join(paragraphs).strip()
        except ValidationException:
            raise
        except (BadZipFile, ValueError, KeyError) as exc:
            raise ValidationException("Invalid DOCX file") from exc
        if not text:
            raise ValidationException("Document does not contain text")
        return text
